import logging
import os
import asyncio
from aiogram import Bot, Dispatcher, types, F
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.filters.command import Command
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton, FSInputFile
from aiogram.utils.keyboard import InlineKeyboardBuilder
from fpdf import FPDF
from docx import Document
from dotenv import load_dotenv
import sys
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
try:
    # Get the directory where the script is located
    script_dir = Path(__file__).parent.absolute()
    env_path = script_dir / '.env'
    
    # Check if .env file exists
    if env_path.exists():
        load_dotenv(dotenv_path=env_path)
        logger.info(f"Loaded environment variables from {env_path}")
    else:
        logger.warning(f".env file not found at {env_path}")
        
    # Get the token from environment variables
    API_TOKEN = os.getenv("API_TOKEN")
    
    # If token is not set, check for direct input
    if not API_TOKEN:
        logger.warning("API_TOKEN not found in environment variables")
        # You can set a default token here if needed
        API_TOKEN = "x"  # Replace with your actual token
        
    logger.info("Bot initialization started")
    
    # Initialize bot and dispatcher
    bot = Bot(token=API_TOKEN)
    storage = MemoryStorage()
    dp = Dispatcher(storage=storage)
    
except Exception as e:
    logger.error(f"Error during initialization: {e}")
    sys.exit(1)

# Define resume templates
TEMPLATES = {
    "professional": "Professional template with a clean, minimalist design",
    "creative": "Creative template with modern styling and accents",
    "academic": "Academic template focused on publications and research",
    "technical": "Technical template highlighting skills and projects"
}

# Resume form states
class ResumeForm(StatesGroup):
    name = State()
    email = State()
    phone = State()
    education = State()
    experience = State()
    skills = State()
    projects = State()
    template = State()
    format = State()

# Start command handler
@dp.message(Command('start'))
async def cmd_start(message: types.Message):
    await message.answer(
        "👋 Welcome to the Resume Builder Bot! I'll help you create a professional resume.\n\n"
        "Let's get started with your information. Type /build to begin creating your resume."
    )

# Build command handler
@dp.message(Command('build'))
async def cmd_build(message: types.Message, state: FSMContext):
    await state.set_state(ResumeForm.name)
    await message.answer("What's your full name?")

# Name handler
@dp.message(ResumeForm.name)
async def process_name(message: types.Message, state: FSMContext):
    await state.update_data(name=message.text)
    await state.set_state(ResumeForm.email)
    await message.answer("Great! What's your email address?")

# Email handler
@dp.message(ResumeForm.email)
async def process_email(message: types.Message, state: FSMContext):
    await state.update_data(email=message.text)
    await state.set_state(ResumeForm.phone)
    await message.answer("What's your phone number?")

# Phone handler
@dp.message(ResumeForm.phone)
async def process_phone(message: types.Message, state: FSMContext):
    await state.update_data(phone=message.text)
    await state.set_state(ResumeForm.education)
    await message.answer(
        "Please enter your education details in this format:\n\n"
        "Degree, Institution, Year\n"
        "Example: 'Bachelor of Science in Computer Science, MIT, 2020'\n\n"
        "You can add multiple entries separated by a new line."
    )

# Education handler
@dp.message(ResumeForm.education)
async def process_education(message: types.Message, state: FSMContext):
    await state.update_data(education=message.text.split('\n'))
    await state.set_state(ResumeForm.experience)
    await message.answer(
        "Please share your work experience in this format:\n\n"
        "Position, Company, Duration, Description\n"
        "Example: 'Software Engineer, Google, 2020-2022, Developed features for Google Maps'\n\n"
        "You can add multiple entries separated by a new line."
    )

# Experience handler
@dp.message(ResumeForm.experience)
async def process_experience(message: types.Message, state: FSMContext):
    await state.update_data(experience=message.text.split('\n'))
    await state.set_state(ResumeForm.skills)
    await message.answer(
        "What are your key skills? Please list them separated by commas.\n"
        "Example: 'Python, JavaScript, Data Analysis, Project Management'"
    )

# Skills handler
@dp.message(ResumeForm.skills)
async def process_skills(message: types.Message, state: FSMContext):
    skills = [skill.strip() for skill in message.text.split(',')]
    await state.update_data(skills=skills)
    await state.set_state(ResumeForm.projects)
    await message.answer(
        "Please list any notable projects (optional). Format:\n\n"
        "Project Name, Description\n"
        "Example: 'Personal Website, Developed a portfolio website using React'\n\n"
        "You can add multiple entries separated by a new line or type 'skip' to continue."
    )

# Projects handler
@dp.message(ResumeForm.projects)
async def process_projects(message: types.Message, state: FSMContext):
    if message.text.lower() == 'skip':
        await state.update_data(projects=[])
    else:
        await state.update_data(projects=message.text.split('\n'))
    
    # Create template selection keyboard
    builder = InlineKeyboardBuilder()
    for template_id, description in TEMPLATES.items():
        builder.add(InlineKeyboardButton(
            text=template_id.capitalize(),
            callback_data=f"template:{template_id}"
        ))
    builder.adjust(2)
    
    await state.set_state(ResumeForm.template)
    await message.answer(
        "Please select a resume template:",
        reply_markup=builder.as_markup()
    )

# Template callback handler
@dp.callback_query(lambda c: c.data.startswith('template:'), ResumeForm.template)
async def process_template(callback_query: types.CallbackQuery, state: FSMContext):
    await callback_query.answer()
    template_id = callback_query.data.split(':')[1]
    
    await state.update_data(template=template_id)
    
    # Create format selection keyboard
    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(text="PDF", callback_data="format:pdf"),
        InlineKeyboardButton(text="DOCX", callback_data="format:docx")
    )
    
    await state.set_state(ResumeForm.format)
    await callback_query.message.answer(
        "Please select the output format:",
        reply_markup=builder.as_markup()
    )

# Format callback handler
@dp.callback_query(lambda c: c.data.startswith('format:'), ResumeForm.format)
async def process_format(callback_query: types.CallbackQuery, state: FSMContext):
    await callback_query.answer()
    format_type = callback_query.data.split(':')[1]
    
    await state.update_data(format=format_type)
    
    # Send processing message
    await callback_query.message.answer(
        "⏳ Processing your resume... This may take a moment."
    )
    
    # Get all data
    resume_data = await state.get_data()
    
    try:
        # Generate manual summary and enhanced skills (without OpenAI)
        summary = generate_summary(resume_data)
        resume_data['summary'] = summary
        
        enhanced_skills = enhance_skills(resume_data['skills'])
        resume_data['enhanced_skills'] = enhanced_skills
        
        # Generate resume
        if resume_data['format'] == 'pdf':
            file_path = generate_pdf_resume(resume_data)
        else:
            file_path = generate_docx_resume(resume_data)
        
        # Send the file
        document = FSInputFile(file_path)
        await callback_query.message.answer_document(
            document=document,
            caption="🎉 Here's your resume!"
        )
        
        # Cleanup
        os.remove(file_path)
        
        # Suggest job recommendations
        jobs = get_job_recommendations(resume_data['skills'])
        job_message = "💼 Based on your skills, here are some job recommendations:\n\n"
        for job in jobs[:5]:
            job_message += f"• {job}\n"
        
        await callback_query.message.answer(job_message)
        
    except Exception as e:
        logging.error(f"Error generating resume: {e}")
        await callback_query.message.answer(
            f"Sorry, there was an error generating your resume: {str(e)}\nPlease try again."
        )
    
    # Reset state
    await state.clear()
    
    await callback_query.message.answer(
        "Type /build to create another resume or /start to see the welcome message again."
    )

# Summary generation (without OpenAI)
def generate_summary(resume_data):
    name = resume_data['name']
    skills = resume_data['skills'][:3]  # Take just the first 3 skills
    
    return f"{name} is a dedicated professional with expertise in {', '.join(skills)}. Seeking new opportunities to apply these skills in a challenging environment."

# Skills enhancement (without OpenAI)
def enhance_skills(skills):
    # Add some common professional skills if less than 8 skills provided
    common_skills = ["Problem Solving", "Communication", "Teamwork", "Time Management", 
                     "Leadership", "Critical Thinking", "Adaptability", "Project Management"]
    
    enhanced = skills.copy()
    
    # Add some common skills if they don't already have many
    if len(enhanced) < 8:
        for skill in common_skills:
            if skill not in enhanced and len(enhanced) < 8:
                enhanced.append(skill)
    
    return enhanced

# Job recommendations (without OpenAI)
def get_job_recommendations(skills):
    # Dictionary mapping skills to related job titles
    skill_to_jobs = {
        "python": ["Python Developer", "Software Engineer", "Data Scientist", "Backend Developer"],
        "javascript": ["Frontend Developer", "Web Developer", "Full Stack Developer", "UI Engineer"],
        "java": ["Java Developer", "Software Engineer", "Android Developer", "Backend Developer"],
        "sql": ["Database Administrator", "Data Analyst", "Backend Developer", "SQL Developer"],
        "data analysis": ["Data Analyst", "Business Analyst", "Data Scientist", "Research Analyst"],
        "machine learning": ["Machine Learning Engineer", "Data Scientist", "AI Researcher", "ML Specialist"],
        "project management": ["Project Manager", "Product Manager", "Scrum Master", "Program Manager"],
        "marketing": ["Marketing Specialist", "Digital Marketer", "Content Strategist", "Marketing Manager"],
        "design": ["UI/UX Designer", "Graphic Designer", "Product Designer", "Web Designer"]
    }
    
    # Default job recommendations
    default_jobs = [
        "Project Manager",
        "Software Developer",
        "Business Analyst",
        "Data Specialist",
        "Marketing Coordinator"
    ]
    
    # Find matching jobs based on skills
    recommended_jobs = set()
    for skill in skills:
        skill_lower = skill.lower()
        for key in skill_to_jobs:
            if key in skill_lower:
                recommended_jobs.update(skill_to_jobs[key])
    
    # If no matching jobs found, use defaults
    if not recommended_jobs:
        return default_jobs
    
    # Convert set to list and return
    return list(recommended_jobs)

# PDF generation
def generate_pdf_resume(data):
    template = data['template']
    filename = f"resume_{data['name'].replace(' ', '_').lower()}.pdf"
    
    pdf = FPDF()
    pdf.add_page()
    
    # Set font based on template
    if template == 'professional':
        title_font = 'Helvetica'
        body_font = 'Helvetica'
    elif template == 'creative':
        title_font = 'Helvetica'
        body_font = 'Helvetica'
    elif template == 'academic':
        title_font = 'Times'
        body_font = 'Times'
    else:  # technical
        title_font = 'Courier'
        body_font = 'Courier'
    
    # Header
    pdf.set_font(title_font, 'B', 16)
    pdf.cell(0, 10, data['name'], 0, 1, 'C')
    
    pdf.set_font(body_font, '', 10)
    pdf.cell(0, 5, f"Email: {data['email']} | Phone: {data['phone']}", 0, 1, 'C')
    pdf.ln(5)
    
    # Summary
    pdf.set_font(title_font, 'B', 12)
    pdf.cell(0, 10, 'PROFESSIONAL SUMMARY', 0, 1)
    pdf.set_font(body_font, '', 10)
    pdf.multi_cell(0, 5, data['summary'])
    pdf.ln(5)
    
    # Experience
    pdf.set_font(title_font, 'B', 12)
    pdf.cell(0, 10, 'WORK EXPERIENCE', 0, 1)
    pdf.set_font(body_font, '', 10)
    
    for exp in data['experience']:
        parts = exp.split(',', 3)
        if len(parts) >= 3:
            position, company, duration = parts[0].strip(), parts[1].strip(), parts[2].strip()
            description = parts[3].strip() if len(parts) > 3 else ""
            
            pdf.set_font(body_font, 'B', 10)
            pdf.cell(0, 5, f"{position} - {company}", 0, 1)
            pdf.set_font(body_font, 'I', 10)
            pdf.cell(0, 5, duration, 0, 1)
            pdf.set_font(body_font, '', 10)
            pdf.multi_cell(0, 5, description)
            pdf.ln(3)
    
    # Education
    pdf.set_font(title_font, 'B', 12)
    pdf.cell(0, 10, 'EDUCATION', 0, 1)
    pdf.set_font(body_font, '', 10)
    
    for edu in data['education']:
        parts = edu.split(',')
        if len(parts) >= 2:
            degree, institution = parts[0].strip(), parts[1].strip()
            year = parts[2].strip() if len(parts) > 2 else ""
            
            pdf.set_font(body_font, 'B', 10)
            pdf.cell(0, 5, degree, 0, 1)
            pdf.set_font(body_font, '', 10)
            pdf.cell(0, 5, f"{institution}, {year}", 0, 1)
            pdf.ln(3)
    
    # Skills
    pdf.set_font(title_font, 'B', 12)
    pdf.cell(0, 10, 'SKILLS', 0, 1)
    pdf.set_font(body_font, '', 10)
    
    skills_text = ', '.join(data['enhanced_skills'])
    pdf.multi_cell(0, 5, skills_text)
    pdf.ln(5)
    
    # Projects (if any)
    if data['projects']:
        pdf.set_font(title_font, 'B', 12)
        pdf.cell(0, 10, 'PROJECTS', 0, 1)
        pdf.set_font(body_font, '', 10)
        
        for project in data['projects']:
            parts = project.split(',', 1)
            if len(parts) >= 1:
                name = parts[0].strip()
                description = parts[1].strip() if len(parts) > 1 else ""
                
                pdf.set_font(body_font, 'B', 10)
                pdf.cell(0, 5, name, 0, 1)
                pdf.set_font(body_font, '', 10)
                pdf.multi_cell(0, 5, description)
                pdf.ln(3)
    
    # Save the PDF
    pdf.output(filename)
    return filename

# DOCX generation
def generate_docx_resume(data):
    template = data['template']
    filename = f"resume_{data['name'].replace(' ', '_').lower()}.docx"
    
    doc = Document()
    
    # Title
    doc.add_heading(data['name'], 0)
    
    # Contact info
    contact = doc.add_paragraph()
    contact.add_run(f"Email: {data['email']} | Phone: {data['phone']}")
    
    # Summary
    doc.add_heading('Professional Summary', 1)
    doc.add_paragraph(data['summary'])
    
    # Experience
    doc.add_heading('Work Experience', 1)
    
    for exp in data['experience']:
        parts = exp.split(',', 3)
        if len(parts) >= 3:
            position, company, duration = parts[0].strip(), parts[1].strip(), parts[2].strip()
            description = parts[3].strip() if len(parts) > 3 else ""
            
            p = doc.add_paragraph()
            p.add_run(f"{position} - {company}\n").bold = True
            p.add_run(f"{duration}\n").italic = True
            p.add_run(description)
    
    # Education
    doc.add_heading('Education', 1)
    
    for edu in data['education']:
        parts = edu.split(',')
        if len(parts) >= 2:
            degree, institution = parts[0].strip(), parts[1].strip()
            year = parts[2].strip() if len(parts) > 2 else ""
            
            p = doc.add_paragraph()
            p.add_run(f"{degree}\n").bold = True
            p.add_run(f"{institution}, {year}")
    
    # Skills
    doc.add_heading('Skills', 1)
    skills_text = ', '.join(data['enhanced_skills'])
    doc.add_paragraph(skills_text)
    
    # Projects (if any)
    if data['projects']:
        doc.add_heading('Projects', 1)
        
        for project in data['projects']:
            parts = project.split(',', 1)
            if len(parts) >= 1:
                name = parts[0].strip()
                description = parts[1].strip() if len(parts) > 1 else ""
                
                p = doc.add_paragraph()
                p.add_run(f"{name}\n").bold = True
                p.add_run(description)
    
    # Save the document
    doc.save(filename)
    return filename

# Main function to start the bot
async def main():
    try:
        logger.info("Starting bot polling")
        # Start the bot
        await dp.start_polling(bot)
    except Exception as e:
        logger.error(f"Error during bot polling: {e}")

if __name__ == '__main__':
    asyncio.run(main())